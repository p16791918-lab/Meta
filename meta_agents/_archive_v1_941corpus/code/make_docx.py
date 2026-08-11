#!/usr/bin/env python3
"""Build a PLOS ONE submission-ready .docx from MANUSCRIPT.md + REFERENCES.md.

- Strips embedded figure images (![](...)) but keeps the figure legends.
- Drops the draft-note italic line and the pointer in the References section,
  and inserts the actual numbered reference entries from REFERENCES.md.
- Renders headings, bold/italic inline, and inline `code` (queries) as monospace.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')


def add_runs(paragraph, text):
    """Add text with **bold**, *italic*, `code` inline formatting."""
    text = re.sub(r'\\([*\[\]_])', r'\1', text)   # unescape markdown backslashes
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(8)
        elif part.startswith('*') and part.endswith('*'):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def reference_entries(path='REFERENCES.md'):
    """Return the numbered reference lines (1. …) from REFERENCES.md, in order."""
    out = []
    for line in open(path, encoding='utf-8'):
        m = re.match(r'^(\d+)\.\s+(.*)', line.rstrip('\n'))
        if m:
            # strip trailing italic notes: provenance *(...)* and/or category *[...]*
            txt = re.sub(r'(\s*\*[\(\[][^*]*[\)\]]\*\s*)+$', '', m.group(2))
            out.append((int(m.group(1)), txt))
    return out


def build():
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    lines = open('MANUSCRIPT.md', encoding='utf-8').read().split('\n')
    in_references = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1
        if not line.strip():
            continue
        # skip horizontal rules and embedded images
        if line.strip() == '---' or line.strip().startswith('!['):
            continue
        # skip the draft-note italic line under the title
        if line.startswith('*Target journal'):
            continue
        # headings
        if line.startswith('# '):
            h = doc.add_heading(level=0); add_runs(h, line[2:])
            continue
        if line.startswith('## '):
            title = line[3:]
            if title.strip().lower() == 'references':
                in_references = True
                doc.add_heading('References', level=1)
                for num, txt in reference_entries():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run(f'{num}. ')
                    add_runs(p, txt)
                # skip the rest of the References section body in MANUSCRIPT.md
                while i < len(lines) and not lines[i].startswith('## '):
                    i += 1
                continue
            doc.add_heading(title, level=1)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            continue
        # normal paragraph
        p = doc.add_paragraph()
        add_runs(p, line)

    doc.save('MANUSCRIPT_PLOS_submission.docx')
    print('wrote MANUSCRIPT_PLOS_submission.docx')


if __name__ == '__main__':
    build()
